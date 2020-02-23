import docx
from docx import *
from docx.shared import Pt
from docx.shared import Inches
import re
from flashtext import KeywordProcessor

from extraction import key_extract
from similar import cos_sim
from summarize import getSummarization

document = Document('ecom.docx')

bold_list = []
font_size = []
font_family = []
data = { 
	'heading': '',
	'module' : {
	}
}

actual_data = {
	
}

i = 0
j = 0
k = 0
content_string = ''

def get_bold_list(para):
	for run in para.runs:
		# font size extract
		if run.font.size:
			font_size.append(run.font.size.pt)

		#font family extract
		if run.font.name:
			font_family.append(run.font.name)

		if run.bold:
			bold_list.append(run.text)

		add_to_dict(run)

def add_to_dict(run):
	global i, j, content_string
	if run.text != '':
		if run.font.size == Pt(20) and run.font.name == 'Arial' :
			data['heading'] = run.text
			i=0
		
		if run.font.size == Pt(15) and run.font.name == 'Arial':
			i = i+1
			data["module"] = {i:{"mod_name": run.text}}
			j=0
		
		if run.font.size == Pt(13) and run.font.name == 'Times New Roman' and re.search('[1-9]\.[1-9]',run.text):
			j=j+1
			data["module"][i].update({j:{"subtopic_name": run.text}})
			actual_data.update({j:run.text})
			content_string = ''

		if run.font.size == Pt(12) and run.font.name == 'Times New Roman':
			content_string += run.text
			data["module"][i][j].update({"content": content_string})



for para in document.paragraphs:
	get_bold_list(para)

# print(data)

# Student functionalities:
def challenge1():
	question = input("Q: ")

	for i in range(1,len(data["module"][2])):	
		similarity = cos_sim(question,data["module"][2][i]["subtopic_name"])
		print(similarity)
		if(similarity > 0.2):
			print(data["module"][2][i]["subtopic_name"])
			ans1 = getSummarization(data["module"][2][i]["content"],5)
			print(ans1)


# Teacher functionalities:
def challenge2():
	doc = docx.Document()
	# print(data["heading"]) #20
	doc.add_heading(data["heading"], 0)
	#print(data["module"][2]["mod_name"]) #18
	doc.add_heading(data["module"][2]["mod_name"])
	for i in range(1,len(data["module"][2])):	
		#print(data["module"][2][i]["subtopic_name"])
		doc.add_heading(data["module"][2][i]["subtopic_name"], level=2)
		#print("\n")
		answer = getSummarization(data["module"][2][i]["content"],1)
		#print(answer)
		paragraph = doc.add_paragraph(answer)
		doc.add_page_break()
		#print("\n\n\n")
	doc.save("demo.docx")
	print("done")